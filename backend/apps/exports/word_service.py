"""
Word 导出服务
使用 python-docx 生成项目报告、知识产权申请报告等 Word 文档
所有导出接口直接返回文件流 HttpResponse
"""
import io

from django.http import HttpResponse


def _doc_to_response(doc, filename):
    """
    Document 转 HttpResponse
    :param doc: docx.Document
    :param filename: 下载文件名（不含扩展名）
    """
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{filename}.docx"
    return response


def _add_kv_table(doc, rows):
    """
    添加键值对表格（两列：字段名、字段值）
    :param doc: docx.Document
    :param rows: [(字段名, 字段值), ...]
    """
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = str(key)
        table.cell(idx, 1).text = str(value)
    return table


class WordExportService:
    """Word 导出服务"""

    @staticmethod
    def export_project_report(project_id):
        """
        导出单项目报告 Word
        包含：基础信息、经费信息、任务信息、贡献记录、知识产权情况
        :param project_id: 项目ID
        """
        from docx import Document
        from docx.shared import Pt
        from apps.projects.models import Project
        from apps.finance.models import FinanceBudget, FinanceExpense
        from apps.tasks.models import Task
        from apps.contributions.models import Contribution
        from apps.intellectual_property.models import IntellectualPropertyApplication

        try:
            project = Project.objects.select_related('leader').get(id=project_id)
        except Project.DoesNotExist:
            raise ValueError('项目不存在')

        doc = Document()
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)

        # 标题
        doc.add_heading(f'项目报告：{project.name}', 0)

        # 一、基础信息
        doc.add_heading('一、基础信息', 1)
        _add_kv_table(doc, [
            ('项目名称', project.name),
            ('项目编号', project.code),
            ('项目负责人', project.leader.name if project.leader else '未指定'),
            ('当前阶段', project.get_current_stage_display()),
            ('项目状态', project.get_status_display()),
            ('开始时间', project.start_date.strftime('%Y-%m-%d') if project.start_date else '未指定'),
            ('预计结束', project.planned_end_date.strftime('%Y-%m-%d') if project.planned_end_date else '未指定'),
            ('项目描述', project.intro or '无'),
        ])

        # 二、经费信息
        doc.add_heading('二、经费信息', 1)
        budget = FinanceBudget.objects.filter(project=project).first()
        if budget:
            _add_kv_table(doc, [
                ('奖金总额', f'{budget.bonus_amount} 元'),
                ('其他收入', f'{budget.other_income} 元'),
                ('已用金额', f'{budget.used_amount} 元'),
                ('待报销', f'{budget.pending_reimbursement} 元'),
                ('剩余金额', f'{budget.remaining_amount} 元'),
                ('经费状态', budget.get_status_display()),
                ('统计周期', budget.period),
            ])
        else:
            doc.add_paragraph('暂无经费预算信息。')

        # 经费明细
        doc.add_heading('经费明细', 2)
        expenses = FinanceExpense.objects.filter(project=project).order_by('-expense_date')
        if expenses:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '支出标题', '金额(元)', '经办人', '支出日期'
            for e in expenses:
                row = table.add_row().cells
                row[0].text = e.title
                row[1].text = str(e.amount)
                row[2].text = e.spender.name if e.spender else ''
                row[3].text = e.expense_date.strftime('%Y-%m-%d') if e.expense_date else ''
        else:
            doc.add_paragraph('暂无经费明细。')

        # 三、任务信息
        doc.add_heading('三、任务信息', 1)
        tasks = Task.objects.filter(project=project).order_by('deadline')
        if tasks:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '任务标题', '指派给', '状态', '截止时间'
            for t in tasks:
                row = table.add_row().cells
                row[0].text = t.title
                row[1].text = t.assignee.name if t.assignee else ''
                row[2].text = t.get_status_display()
                row[3].text = t.deadline.strftime('%Y-%m-%d') if t.deadline else ''
        else:
            doc.add_paragraph('暂无任务记录。')

        # 四、贡献记录
        doc.add_heading('四、贡献记录', 1)
        contributions = Contribution.objects.filter(project=project).order_by('-created_at')
        if contributions:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '贡献人', '贡献类型', '权重', '审核状态'
            for c in contributions:
                row = table.add_row().cells
                row[0].text = c.user.name if c.user else ''
                row[1].text = c.get_contribution_type_display()
                row[2].text = str(c.weight)
                row[3].text = c.get_status_display()
        else:
            doc.add_paragraph('暂无贡献记录。')

        # 五、知识产权情况
        doc.add_heading('五、知识产权情况', 1)
        ip_apps = IntellectualPropertyApplication.objects.filter(related_project=project)
        if ip_apps:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = '成果名称', '成果类型', '当前状态'
            for a in ip_apps:
                row = table.add_row().cells
                row[0].text = a.title
                row[1].text = a.get_ip_type_display()
                row[2].text = a.get_status_display()
        else:
            doc.add_paragraph('暂无知识产权申请。')

        # 页脚
        doc.add_paragraph()
        doc.add_paragraph(f'报告生成时间：{__import__("django").utils.timezone.now().strftime("%Y-%m-%d %H:%M")}')

        return _doc_to_response(doc, f'项目报告_{project.name}')

    @staticmethod
    def export_ip_report(ip_id):
        """
        导出知识产权申请报告 Word
        :param ip_id: 知识产权申请ID
        """
        from docx import Document
        from docx.shared import Pt
        from apps.intellectual_property.models import IntellectualPropertyApplication

        try:
            application = IntellectualPropertyApplication.objects.select_related(
                'related_project', 'main_writer', 'applicant_executor'
            ).get(id=ip_id)
        except IntellectualPropertyApplication.DoesNotExist:
            raise ValueError('知识产权申请不存在')

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)

        doc.add_heading(f'知识产权申请报告：{application.title}', 0)

        # 一、基本信息
        doc.add_heading('一、基本信息', 1)
        _add_kv_table(doc, [
            ('成果名称', application.title),
            ('内部编号', application.application_code),
            ('成果类型', application.get_ip_type_display()),
            ('关联项目', application.related_project.name if application.related_project else '无'),
            ('当前状态', application.get_status_display()),
        ])

        # 二、人员信息
        doc.add_heading('二、人员信息', 1)
        _add_kv_table(doc, [
            ('主导撰写人', application.main_writer.name if application.main_writer else '未指定'),
            ('申请执行人', application.applicant_executor.name if application.applicant_executor else '未指定'),
            ('退回修改次数', str(application.return_count)),
        ])

        # 三、关键日期
        doc.add_heading('三、关键日期', 1)
        _add_kv_table(doc, [
            ('提交日期', application.submit_date.strftime('%Y-%m-%d') if application.submit_date else '未提交'),
            ('受理日期', application.accepted_date.strftime('%Y-%m-%d') if application.accepted_date else '未受理'),
            ('授权日期', application.authorized_date.strftime('%Y-%m-%d') if application.authorized_date else '未授权'),
        ])

        # 四、流程记录（退回修改记录）
        doc.add_heading('四、流程记录', 1)
        return_records = application.return_records.all().order_by('-return_time')
        if return_records:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '退回时间', '退回来源', '退回原因', '处理结果'
            for record in return_records:
                row = table.add_row().cells
                row[0].text = record.return_time.strftime('%Y-%m-%d') if record.return_time else ''
                row[1].text = record.get_return_source_display()
                row[2].text = record.return_reason
                row[3].text = record.get_result_display()
        else:
            doc.add_paragraph('暂无流程记录。')

        doc.add_paragraph()
        doc.add_paragraph(f'报告生成时间：{__import__("django").utils.timezone.now().strftime("%Y-%m-%d %H:%M")}')

        return _doc_to_response(doc, f'知识产权报告_{application.title}')
