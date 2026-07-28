"""定时报表的时间计算、文件生成、邮件投递与执行编排。"""
from __future__ import annotations

import calendar
import io
import logging
import re
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from django.conf import settings
from django.core.files.base import ContentFile
from django.core.mail import EmailMessage
from django.db import transaction
from django.utils import timezone

from apps.notifications.models import Notification
from .scheduled_report_models import ScheduledReport, ScheduledReportExecution

logger = logging.getLogger(__name__)

DEFAULT_SCHEDULE_RETRY_SECONDS = 5 * 60
DEFAULT_SCHEDULE_EXECUTION_TIMEOUT_SECONDS = 30 * 60


def compute_next_run(schedule: ScheduledReport, base=None):
    """根据计划的本地时区、频率与执行时刻计算下一次 UTC 时间。"""
    base = base or timezone.now()
    try:
        local_timezone = ZoneInfo(schedule.timezone or settings.TIME_ZONE)
    except ZoneInfoNotFoundError:
        local_timezone = ZoneInfo(settings.TIME_ZONE)

    local_base = timezone.localtime(base, local_timezone)
    candidate = datetime.combine(
        local_base.date(),
        schedule.execution_time,
        tzinfo=local_timezone,
    )

    if schedule.frequency == ScheduledReport.Frequency.DAILY:
        if candidate <= local_base:
            candidate += timedelta(days=1)
    elif schedule.frequency == ScheduledReport.Frequency.WEEKLY:
        days_ahead = (min(max(schedule.weekday, 0), 6) - local_base.weekday()) % 7
        candidate += timedelta(days=days_ahead)
        if candidate <= local_base:
            candidate += timedelta(weeks=1)
    else:
        day = min(max(schedule.day_of_month, 1), 28)
        candidate = candidate.replace(day=day)
        if candidate <= local_base:
            year = local_base.year + (1 if local_base.month == 12 else 0)
            month = 1 if local_base.month == 12 else local_base.month + 1
            safe_day = min(day, calendar.monthrange(year, month)[1])
            candidate = candidate.replace(year=year, month=month, day=safe_day)

    return candidate.astimezone(ZoneInfo('UTC'))


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n]+', '_', value).strip(' ._')
    return (cleaned or 'report')[:100]


def _report_rows(data: dict) -> tuple[list[tuple[str, object]], list[dict]]:
    summary_labels = {
        'total': '总数',
        'done': '已完成',
        'overdue': '已逾期',
        'doing': '进行中',
        'todo': '待处理',
        'active': '进行中项目',
        'closed': '已结项项目',
        'paused': '已暂停项目',
        'awarded': '已获奖',
        'promoted': '已晋级',
        'total_amount': '支出总额',
        'count': '记录数',
        'message': '说明',
    }
    detail_labels = {
        'key': '编码',
        'label': '分组',
        'count': '数量',
        'total': '金额',
    }
    summary = [
        (summary_labels.get(str(key), str(key)), value)
        for key, value in (data.get('summary') or {}).items()
    ]
    groups = [
        {detail_labels.get(str(key), str(key)): value for key, value in row.items()}
        for row in (data.get('groups') or [])
    ]
    return summary, groups


def _generate_xlsx(name: str, data: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = '概览'
    summary_sheet.append(['报表', name])
    summary_sheet.append(['生成时间', timezone.localtime().strftime('%Y-%m-%d %H:%M:%S')])
    summary_sheet.append([])
    summary_sheet.append(['指标', '数值'])
    for cell in summary_sheet[4]:
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = PatternFill('solid', fgColor='176B73')
    summary, groups = _report_rows(data)
    for label, value in summary:
        summary_sheet.append([label, value])
    summary_sheet.column_dimensions['A'].width = 24
    summary_sheet.column_dimensions['B'].width = 28

    detail_sheet = workbook.create_sheet('明细')
    if groups:
        headers = list(dict.fromkeys(key for row in groups for key in row.keys()))
        detail_sheet.append(headers)
        for cell in detail_sheet[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill('solid', fgColor='176B73')
        for row in groups:
            detail_sheet.append([row.get(header, '') for header in headers])
        for index in range(1, len(headers) + 1):
            detail_sheet.column_dimensions[chr(64 + index)].width = 20
    else:
        detail_sheet.append(['说明'])
        detail_sheet.append(['当前筛选条件下暂无分组数据'])

    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _generate_docx(name: str, data: dict) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(name, level=0)
    document.add_paragraph(f'生成时间：{timezone.localtime():%Y-%m-%d %H:%M:%S}')
    document.add_heading('概览', level=1)
    summary, groups = _report_rows(data)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.rows[0].cells[0].text = '指标'
    summary_table.rows[0].cells[1].text = '数值'
    for label, value in summary:
        cells = summary_table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)

    document.add_heading('明细', level=1)
    if groups:
        headers = list(dict.fromkeys(key for row in groups for key in row.keys()))
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in groups:
            cells = table.add_row().cells
            for index, header in enumerate(headers):
                cells[index].text = str(row.get(header, ''))
    else:
        document.add_paragraph('当前筛选条件下暂无分组数据。')

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _generate_pdf(name: str, data: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    # ReportLab 内置的 CID 字体不依赖系统 GTK/Pango，在 Windows 与容器中
    # 都能稳定生成中文 PDF。
    font_name = 'STSong-Light'
    try:
        pdfmetrics.getFont(font_name)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title=name,
        author='团队管理平台',
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Title'],
        fontName=font_name,
        fontSize=20,
        leading=28,
        textColor=colors.HexColor('#176B73'),
        alignment=TA_LEFT,
        spaceAfter=5 * mm,
    )
    heading_style = ParagraphStyle(
        'ReportHeading',
        parent=styles['Heading2'],
        fontName=font_name,
        fontSize=13,
        leading=18,
        textColor=colors.HexColor('#26332F'),
        spaceBefore=4 * mm,
        spaceAfter=3 * mm,
    )
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['BodyText'],
        fontName=font_name,
        fontSize=9.5,
        leading=14,
        textColor=colors.HexColor('#4C5B56'),
        alignment=TA_LEFT,
    )
    header_style = ParagraphStyle(
        'ReportTableHeader',
        parent=body_style,
        textColor=colors.white,
        alignment=TA_CENTER,
    )
    cell_style = ParagraphStyle(
        'ReportTableCell',
        parent=body_style,
        textColor=colors.HexColor('#26332F'),
    )

    def paragraph(value, style=cell_style):
        text = str(value).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return Paragraph(text, style)

    def decorate_table(table):
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#176B73')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [
                colors.white,
                colors.HexColor('#F4F8F7'),
            ]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D8E0DD')),
            ('LEFTPADDING', (0, 0), (-1, -1), 7),
            ('RIGHTPADDING', (0, 0), (-1, -1), 7),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        return table

    summary, groups = _report_rows(data)
    story = [
        Paragraph(str(name), title_style),
        Paragraph(
            f'生成时间：{timezone.localtime():%Y-%m-%d %H:%M:%S}',
            body_style,
        ),
        Spacer(1, 3 * mm),
        Paragraph('概览', heading_style),
    ]
    summary_data = [
        [Paragraph('指标', header_style), Paragraph('数值', header_style)],
        *[[paragraph(label), paragraph(value)] for label, value in summary],
    ]
    story.append(decorate_table(Table(
        summary_data,
        colWidths=[92 * mm, 76 * mm],
        repeatRows=1,
    )))
    story.append(Paragraph('明细', heading_style))

    if groups:
        headers = list(dict.fromkeys(key for row in groups for key in row.keys()))
        detail_data = [
            [Paragraph(str(header), header_style) for header in headers],
            *[
                [paragraph(row.get(header, '')) for header in headers]
                for row in groups
            ],
        ]
        story.append(decorate_table(Table(
            detail_data,
            colWidths=[168 * mm / len(headers)] * len(headers),
            repeatRows=1,
        )))
    else:
        story.append(Paragraph('当前筛选条件下暂无分组数据。', body_style))

    def draw_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor('#176B73'))
        canvas.setLineWidth(1.2)
        canvas.line(18 * mm, A4[1] - 12 * mm, A4[0] - 18 * mm, A4[1] - 12 * mm)
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor('#72807B'))
        canvas.drawRightString(A4[0] - 18 * mm, 9 * mm, f'第 {doc.page} 页')
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return output.getvalue()


def _is_internal_user(user) -> bool:
    return bool(
        user
        and user.is_active
        and user.membership_status in {'active', 'on_leave'}
    )


def _schedule_creator(schedule: ScheduledReport):
    """兼容迁移前未写入 created_by 的计划，以关联报表所有者作为创建人。"""
    return schedule.created_by or schedule.report.created_by


def _recipient_snapshot(schedule: ScheduledReport) -> list[dict]:
    return [
        {'id': item.pk, 'name': item.name, 'email': item.email}
        for item in schedule.recipients.all()
        if _is_internal_user(item)
    ]


def report_visible_project_ids(report, user) -> set[int]:
    """返回报表按当前筛选条件会读取到的项目集合。"""
    from common.project_access import scope_project_queryset

    config = report.config or {}
    data_source = config.get('data_source', 'project')
    filters = config.get('filters', {}) or {}

    if data_source == 'project':
        from apps.projects.models import Project

        queryset = scope_project_queryset(
            Project.objects.all(),
            user,
            project_lookup='',
        )
        if filters.get('status'):
            queryset = queryset.filter(status=filters['status'])
        return set(queryset.values_list('id', flat=True))

    model = None
    if data_source == 'task':
        from apps.tasks.models import Task

        model = Task
    elif data_source == 'finance':
        from apps.finance.models import FinanceExpense

        model = FinanceExpense
    elif data_source == 'competition':
        from apps.competitions.models import Competition

        model = Competition
    if model is None:
        return set()

    queryset = scope_project_queryset(
        model.objects.all(),
        user,
        project_lookup='project',
    )
    if filters.get('project_id'):
        queryset = queryset.filter(project_id=filters['project_id'])
    if data_source == 'task' and filters.get('status'):
        queryset = queryset.filter(status=filters['status'])
    elif data_source == 'finance' and filters.get('category'):
        queryset = queryset.filter(category=filters['category'])
    elif data_source == 'competition' and filters.get('level'):
        queryset = queryset.filter(level=filters['level'])
    return set(queryset.values_list('project_id', flat=True))


def report_recipient_scope_error(report, creator, recipients) -> str:
    """确保接收人既属于允许的组织范围，也能看到报表涉及的每个项目。"""
    from apps.projects.models import Project
    from apps.users.models import User
    from common.project_access import scope_organization_users, scope_project_queryset

    recipients = list(recipients)
    if any(not _is_internal_user(recipient) for recipient in recipients):
        return '接收人必须是在队或暂离的内部成员'

    recipient_ids = {recipient.pk for recipient in recipients}
    organization_recipient_ids = set(
        scope_organization_users(
            User.objects.filter(pk__in=recipient_ids),
            creator,
        ).values_list('pk', flat=True)
    )
    if recipient_ids - organization_recipient_ids:
        return '接收人必须与计划创建人属于同一团队组织'

    report_project_ids = report_visible_project_ids(report, creator)
    if not report_project_ids:
        return ''
    for recipient in recipients:
        visible_project_ids = set(
            scope_project_queryset(
                Project.objects.filter(pk__in=report_project_ids),
                recipient,
                project_lookup='',
            ).values_list('pk', flat=True)
        )
        if report_project_ids - visible_project_ids:
            return f'接收人“{recipient.name}”无权查看报表涉及的全部项目'
    return ''


def _positive_seconds(setting_name: str, default: int) -> int:
    try:
        value = int(getattr(settings, setting_name, default))
    except (TypeError, ValueError):
        return default
    return value if value > 0 else default


def _retry_at(now):
    seconds = _positive_seconds(
        'SCHEDULED_REPORT_RETRY_SECONDS',
        DEFAULT_SCHEDULE_RETRY_SECONDS,
    )
    return now + timedelta(seconds=seconds)


def schedule_scope_error(schedule: ScheduledReport) -> str:
    """执行前再次校验资源归属和接收人数据域，防止旧数据或后台写入绕过 API。"""
    creator = _schedule_creator(schedule)
    if not _is_internal_user(creator):
        return '计划创建人不是有效内部成员'
    if (
        creator.global_role not in ('teacher', 'sys_admin')
        and schedule.report.created_by_id != creator.id
    ):
        return '计划引用了创建人无权使用的报表'
    if any(not _is_internal_user(user) for user in schedule.recipients.all()):
        return '计划包含外部、离队或已停用接收人'
    recipient_error = report_recipient_scope_error(
        schedule.report,
        creator,
        schedule.recipients.all(),
    )
    if recipient_error:
        return recipient_error
    return ''


def generate_report_file(schedule: ScheduledReport) -> tuple[str, bytes]:
    """生成所选格式的报表文件。"""
    from .custom_report_views import _generate_report_data

    data = _generate_report_data(schedule.report, user=_schedule_creator(schedule))
    file_name = (
        f'{_safe_filename(schedule.report.name)}_'
        f'{timezone.localtime():%Y%m%d_%H%M%S}.{schedule.file_format}'
    )
    generators = {
        ScheduledReport.FileFormat.XLSX: _generate_xlsx,
        ScheduledReport.FileFormat.DOCX: _generate_docx,
        ScheduledReport.FileFormat.PDF: _generate_pdf,
    }
    return file_name, generators[schedule.file_format](schedule.report.name, data)


def _deliver_email(schedule: ScheduledReport, file_name: str, content: bytes) -> tuple[str, str]:
    from apps.notifications.services import should_notify_user

    configured_users = [
        user for user in schedule.recipients.all() if _is_internal_user(user)
    ]
    if not configured_users:
        return ScheduledReportExecution.DeliveryStatus.NOT_REQUESTED, '未配置接收人，仅保留站内文件'
    recipients = [
        user.email
        for user in configured_users
        if user.email and should_notify_user(
            user,
            category=Notification.NotificationType.REPORT,
            channel=Notification.Channel.EMAIL,
            priority=Notification.Priority.NORMAL,
        )
    ]
    if not recipients:
        return ScheduledReportExecution.DeliveryStatus.NOT_REQUESTED, '接收人已关闭报表邮件或当前处于免打扰时段'
    if not settings.EMAIL_HOST_USER:
        return ScheduledReportExecution.DeliveryStatus.NOT_CONFIGURED, '邮件服务未配置，仅保留站内文件'

    try:
        message = EmailMessage(
            subject=f'定时报表：{schedule.report.name}',
            body='报表已按计划生成，文件见附件。',
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=recipients,
        )
        message.attach(file_name, content)
        sent_count = message.send(fail_silently=False)
        if sent_count != 1:
            return ScheduledReportExecution.DeliveryStatus.FAILED, '邮件后端未确认投递'
        return ScheduledReportExecution.DeliveryStatus.SENT, f'已发送至 {len(recipients)} 个邮箱'
    except Exception as exc:  # pragma: no cover - 取决于外部邮件服务
        logger.exception('定时报表邮件发送失败')
        return ScheduledReportExecution.DeliveryStatus.FAILED, f'邮件发送失败：{exc}'


def _notify(schedule: ScheduledReport, execution: ScheduledReportExecution) -> None:
    from apps.notifications.services import NotificationService

    creator = _schedule_creator(schedule)
    users = [
        user for user in schedule.recipients.all() if _is_internal_user(user)
    ]
    if (
        _is_internal_user(creator)
        and all(user.pk != creator.pk for user in users)
    ):
        users.append(creator)
    title = '定时报表已生成' if execution.status != ScheduledReport.RunStatus.FAILED else '定时报表生成失败'
    NotificationService.bulk_create_notifications(
        recipients=users,
        title=title,
        content=f'{schedule.report.name}：{execution.message or execution.error}',
        category=Notification.NotificationType.REPORT,
        sender=execution.generated_by,
        priority=(
            Notification.Priority.NORMAL
            if execution.status != ScheduledReport.RunStatus.FAILED
            else Notification.Priority.HIGH
        ),
        ref_type='scheduled_report',
        ref_id=schedule.pk,
    )


def execute_scheduled_report(
    schedule: ScheduledReport,
    *,
    trigger=ScheduledReportExecution.Trigger.MANUAL,
    user=None,
    execution: ScheduledReportExecution | None = None,
) -> ScheduledReportExecution:
    """执行一次计划，持久化文件、投递结果和站内通知。"""
    if execution is not None and execution.schedule_id != schedule.pk:
        raise ValueError('执行记录与定时报表计划不匹配')
    if execution is not None and execution.status != ScheduledReport.RunStatus.RUNNING:
        # Celery 重复投递同一领取记录时保持幂等，不重复生成和发送。
        return execution
    if execution is None:
        execution = ScheduledReportExecution.objects.create(
            schedule=schedule,
            trigger=trigger,
            generated_by=(
                user if getattr(user, 'is_authenticated', False) else None
            ),
            recipient_snapshot=_recipient_snapshot(schedule),
        )
    else:
        trigger = execution.trigger
    schedule.last_status = ScheduledReport.RunStatus.RUNNING
    schedule.last_error = ''
    schedule.save(update_fields=['last_status', 'last_error'])

    scope_error = schedule_scope_error(schedule)
    if scope_error:
        now = timezone.now()
        execution.status = ScheduledReport.RunStatus.FAILED
        execution.delivery_status = (
            ScheduledReportExecution.DeliveryStatus.NOT_REQUESTED
        )
        execution.error = scope_error
        execution.message = '报表计划权限校验失败，已自动停用'
        execution.finished_at = now
        execution.save()
        schedule.is_active = False
        schedule.next_run = None
        schedule.last_run = now
        schedule.last_status = ScheduledReport.RunStatus.FAILED
        schedule.last_error = scope_error
        schedule.save(update_fields=[
            'is_active', 'next_run', 'last_run', 'last_status', 'last_error',
        ])
        _notify(schedule, execution)
        return execution

    try:
        file_name, content = generate_report_file(schedule)
        execution.file.save(file_name, ContentFile(content), save=False)
        execution.file_name = file_name
        execution.file_format = schedule.file_format
        execution.file_size = len(content)
        delivery_status, delivery_message = _deliver_email(schedule, file_name, content)
        execution.delivery_status = delivery_status
        execution.status = (
            ScheduledReport.RunStatus.PARTIAL
            if delivery_status in {
                ScheduledReportExecution.DeliveryStatus.NOT_CONFIGURED,
                ScheduledReportExecution.DeliveryStatus.FAILED,
            }
            else ScheduledReport.RunStatus.SUCCESS
        )
        execution.message = delivery_message
    except Exception as exc:
        logger.exception('定时报表生成失败')
        execution.status = ScheduledReport.RunStatus.FAILED
        execution.delivery_status = ScheduledReportExecution.DeliveryStatus.NOT_REQUESTED
        execution.error = str(exc)
        execution.message = '报表生成失败'

    now = timezone.now()
    execution.finished_at = now
    execution.save()
    schedule.last_run = now
    if not schedule.is_active:
        schedule.next_run = None
    elif (
        execution.status == ScheduledReport.RunStatus.FAILED
        and trigger == ScheduledReportExecution.Trigger.SCHEDULED
    ):
        schedule.next_run = _retry_at(now)
    else:
        schedule.next_run = compute_next_run(schedule, base=now)
    schedule.last_status = execution.status
    schedule.last_error = execution.error
    schedule.save(update_fields=['last_run', 'next_run', 'last_status', 'last_error'])
    _notify(schedule, execution)
    return execution


def fail_scheduled_report_execution(
    execution: ScheduledReportExecution,
    error,
    *,
    now=None,
) -> ScheduledReportExecution:
    """把 Worker 外层异常写回领取记录，并安排短延迟重试。"""
    now = now or timezone.now()
    error_text = str(error).strip() or error.__class__.__name__
    with transaction.atomic():
        locked_execution = (
            ScheduledReportExecution.objects.select_for_update()
            .select_related('schedule', 'schedule__report')
            .get(pk=execution.pk)
        )
        if locked_execution.status != ScheduledReport.RunStatus.RUNNING:
            return locked_execution

        schedule = ScheduledReport.objects.select_for_update().get(
            pk=locked_execution.schedule_id
        )
        locked_execution.status = ScheduledReport.RunStatus.FAILED
        locked_execution.delivery_status = (
            ScheduledReportExecution.DeliveryStatus.NOT_REQUESTED
        )
        locked_execution.error = error_text
        locked_execution.message = '报表任务执行异常，已安排重试'
        locked_execution.finished_at = now
        locked_execution.save(update_fields=[
            'status', 'delivery_status', 'error', 'message', 'finished_at',
        ])

        schedule.last_run = now
        schedule.next_run = _retry_at(now) if schedule.is_active else None
        schedule.last_status = ScheduledReport.RunStatus.FAILED
        schedule.last_error = error_text
        schedule.save(update_fields=[
            'last_run', 'next_run', 'last_status', 'last_error',
        ])

    _notify(schedule, locked_execution)
    return locked_execution


def claim_due_schedule_execution_ids(now=None, limit=100) -> list[int]:
    """
    原子领取到期计划并预先创建 RUNNING 执行记录。

    next_run 在执行成功前保持到期状态；新鲜 RUNNING 记录负责去重，超时记录
    会被持久化为失败后重新领取，因此 Worker 崩溃不会永久漏跑。
    """
    now = now or timezone.now()
    timeout_seconds = _positive_seconds(
        'SCHEDULED_REPORT_EXECUTION_TIMEOUT_SECONDS',
        DEFAULT_SCHEDULE_EXECUTION_TIMEOUT_SECONDS,
    )
    stale_before = now - timedelta(seconds=timeout_seconds)
    claimed: list[int] = []
    limit = max(1, min(int(limit), 1000))

    with transaction.atomic():
        schedules = list(
            ScheduledReport.objects.select_for_update(skip_locked=True)
            .filter(is_active=True, next_run__lte=now)
            .order_by('next_run')
            [:limit]
        )
        for schedule in schedules:
            running = schedule.executions.filter(
                trigger=ScheduledReportExecution.Trigger.SCHEDULED,
                status=ScheduledReport.RunStatus.RUNNING,
            )
            running.filter(started_at__lte=stale_before).update(
                status=ScheduledReport.RunStatus.FAILED,
                delivery_status=(
                    ScheduledReportExecution.DeliveryStatus.NOT_REQUESTED
                ),
                error='Worker 执行超时，系统已重新领取',
                message='上一次定时报表执行超时',
                finished_at=now,
            )
            if running.filter(started_at__gt=stale_before).exists():
                continue

            execution = ScheduledReportExecution.objects.create(
                schedule=schedule,
                trigger=ScheduledReportExecution.Trigger.SCHEDULED,
                recipient_snapshot=_recipient_snapshot(schedule),
            )
            claimed.append(execution.pk)
            schedule.last_status = ScheduledReport.RunStatus.RUNNING
            schedule.last_error = ''
            schedule.save(update_fields=['last_status', 'last_error'])
    return claimed


def claim_due_schedule_ids(now=None) -> list[int]:
    """兼容旧调用方：领取仍持久化执行记录，但返回计划 ID。"""
    execution_ids = claim_due_schedule_execution_ids(now=now)
    schedule_by_execution = dict(
        ScheduledReportExecution.objects.filter(pk__in=execution_ids)
        .values_list('pk', 'schedule_id')
    )
    return [
        schedule_by_execution[execution_id]
        for execution_id in execution_ids
        if execution_id in schedule_by_execution
    ]
