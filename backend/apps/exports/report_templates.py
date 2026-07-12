"""
项目报告模板生成
使用 python-docx 生成项目完整报告（Word 格式）
若 python-docx 未安装，降级为纯文本导出
"""
import io

from django.utils import timezone

# ============ 惰性导入 python-docx，降级处理 ============
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:  # pragma: no cover - 依赖环境
    HAS_PYTHON_DOCX = False


def _gather_project_data(project_id):
    """
    汇总项目报告所需的全部数据
    :param project_id: 项目ID
    :return: dict 包含项目及关联数据
    :raises ValueError: 项目不存在
    """
    from apps.projects.models import Project, ProjectMember, ProjectStageLog
    from apps.competitions.models import Competition
    from apps.finance.models import FinanceBudget, FinanceExpense
    from apps.intellectual_property.models import IntellectualPropertyApplication
    from apps.contributions.models import Contribution

    try:
        project = Project.objects.select_related('leader').get(id=project_id)
    except Project.DoesNotExist:
        raise ValueError('项目不存在')

    # 阶段历程（按时间正序）
    stage_logs = list(
        ProjectStageLog.objects.filter(project=project)
        .select_related('operator')
        .order_by('created_at')
    )

    # 比赛记录
    competitions = list(
        Competition.objects.filter(project=project).order_by('-created_at')
    )

    # 经费统计
    budget = FinanceBudget.objects.filter(project=project).first()
    expenses = list(
        FinanceExpense.objects.filter(project=project)
        .select_related('spender', 'reviewer')
        .order_by('-expense_date')
    )

    # 项目成员
    members = list(
        ProjectMember.objects.filter(project=project)
        .select_related('user')
        .order_by('joined_at')
    )

    # 知识产权
    ip_apps = list(
        IntellectualPropertyApplication.objects.filter(related_project=project)
        .select_related('main_writer', 'applicant_executor')
        .order_by('-created_at')
    )

    # 贡献记录
    contributions = list(
        Contribution.objects.filter(project=project)
        .select_related('user', 'reviewer', 'filled_by')
        .order_by('-created_at')
    )

    return {
        'project': project,
        'stage_logs': stage_logs,
        'competitions': competitions,
        'budget': budget,
        'expenses': expenses,
        'members': members,
        'ip_apps': ip_apps,
        'contributions': contributions,
    }


def _fmt_date(value):
    """格式化日期为 YYYY-MM-DD，None 返回空串"""
    if value:
        return value.strftime('%Y-%m-%d')
    return ''


def _fmt_datetime(value):
    """格式化日期时间为 YYYY-MM-DD HH:MM，None 返回空串"""
    if value:
        return value.strftime('%Y-%m-%d %H:%M')
    return ''


# ============================================================
# Word（python-docx）报告生成
# ============================================================

def _add_kv_table(doc, rows):
    """添加键值对表格（两列：字段名、字段值）"""
    table = doc.add_table(rows=len(rows), cols=2)
    try:
        table.style = 'Light Grid Accent 1'
    except Exception:
        pass
    for idx, (key, value) in enumerate(rows):
        table.cell(idx, 0).text = str(key)
        table.cell(idx, 1).text = str(value)
    return table


def _add_data_table(doc, headers, data_rows):
    """添加多列表格（含表头）"""
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Light Grid Accent 1'
    except Exception:
        pass
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row_data in data_rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val) if val is not None else ''
    return table


def _generate_docx_report(data):
    """使用 python-docx 生成 Word 报告，返回 BytesIO"""
    project = data['project']
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    # ========== 封面标题 ==========
    title = doc.add_heading(f'项目报告：{project.name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # 空行

    # ========== 一、项目基本信息 ==========
    doc.add_heading('一、项目基本信息', 1)
    _add_kv_table(doc, [
        ('项目名称', project.name),
        ('项目编号', project.code),
        ('项目负责人', project.leader.name if project.leader else '未指定'),
        ('当前阶段', project.get_current_stage_display()),
        ('项目状态', project.get_status_display()),
        ('优先级', project.get_priority_display()),
        ('开始时间', _fmt_date(project.start_date)),
        ('预计结束', _fmt_date(project.planned_end_date)),
        ('实际结束', _fmt_date(project.actual_end_date)),
        ('项目简介', project.intro or '无'),
        ('创建时间', _fmt_datetime(project.created_at)),
    ])

    # ========== 二、阶段历程 ==========
    doc.add_heading('二、阶段历程', 1)
    stage_logs = data['stage_logs']
    if stage_logs:
        rows = []
        for log in stage_logs:
            from_stage = log.get_from_stage_display() if log.from_stage else '初始'
            to_stage = log.get_to_stage_display()
            operator = log.operator.name if log.operator else '系统'
            rows.append([
                _fmt_datetime(log.created_at),
                from_stage,
                to_stage,
                operator,
                log.note or '',
            ])
        _add_data_table(doc, ['变更时间', '原阶段', '目标阶段', '操作人', '备注'], rows)
    else:
        doc.add_paragraph('暂无阶段变更记录。')

    # ========== 三、比赛记录 ==========
    doc.add_heading('三、比赛记录', 1)
    competitions = data['competitions']
    if competitions:
        rows = []
        for c in competitions:
            rows.append([
                c.name,
                c.get_level_display(),
                c.organizer or '',
                c.get_status_display(),
                '是' if c.is_promoted else '否',
                '是' if c.is_awarded else '否',
                c.award_level or '',
                _fmt_date(c.result_date),
            ])
        _add_data_table(
            doc,
            ['比赛名称', '级别', '主办方', '状态', '是否晋级', '是否获奖', '获奖等级', '结果公布'],
            rows,
        )
    else:
        doc.add_paragraph('暂无比赛记录。')

    # ========== 四、经费统计 ==========
    doc.add_heading('四、经费统计', 1)
    budget = data['budget']
    if budget:
        _add_kv_table(doc, [
            ('奖金总额', f'{budget.bonus_amount} 元'),
            ('其他收入', f'{budget.other_income} 元'),
            ('总收入', f'{budget.total_income} 元'),
            ('已用金额', f'{budget.used_amount} 元'),
            ('待报销', f'{budget.pending_reimbursement} 元'),
            ('剩余金额', f'{budget.remaining_amount} 元'),
            ('经费状态', budget.get_status_display()),
            ('统计周期', budget.period or ''),
        ])
    else:
        doc.add_paragraph('暂无经费预算信息。')

    # 经费明细
    doc.add_heading('经费明细', 2)
    expenses = data['expenses']
    if expenses:
        rows = []
        for e in expenses:
            rows.append([
                e.title,
                str(e.amount),
                e.spender.name if e.spender else '',
                _fmt_date(e.expense_date),
                e.get_category_display(),
                e.purpose or '',
            ])
        _add_data_table(
            doc,
            ['支出标题', '金额(元)', '经办人', '支出日期', '类别', '用途说明'],
            rows,
        )
    else:
        doc.add_paragraph('暂无经费明细。')

    # ========== 五、成员列表 ==========
    doc.add_heading('五、成员列表', 1)
    members = data['members']
    if members:
        rows = []
        for m in members:
            rows.append([
                m.user.name,
                m.user.email,
                m.get_role_in_project_display(),
                _fmt_datetime(m.joined_at),
            ])
        _add_data_table(doc, ['姓名', '邮箱', '项目角色', '加入时间'], rows)
    else:
        doc.add_paragraph('暂无项目成员。')

    # ========== 六、知识产权 ==========
    doc.add_heading('六、知识产权', 1)
    ip_apps = data['ip_apps']
    if ip_apps:
        rows = []
        for a in ip_apps:
            rows.append([
                a.title,
                a.application_code,
                a.get_ip_type_display(),
                a.get_status_display(),
                a.main_writer.name if a.main_writer else '',
                _fmt_date(a.submit_date),
                _fmt_date(a.accepted_date),
                _fmt_date(a.authorized_date),
            ])
        _add_data_table(
            doc,
            ['成果名称', '内部编号', '成果类型', '当前状态', '主导撰写人', '提交日期', '受理日期', '授权日期'],
            rows,
        )
    else:
        doc.add_paragraph('暂无知识产权申请。')

    # ========== 七、贡献汇总 ==========
    doc.add_heading('七、贡献汇总', 1)
    contributions = data['contributions']
    if contributions:
        rows = []
        for c in contributions:
            rows.append([
                c.user.name if c.user else '',
                c.get_contribution_type_display(),
                c.content or c.description or '',
                str(c.weight),
                c.get_status_display(),
                c.period or '',
                _fmt_datetime(c.created_at),
            ])
        _add_data_table(
            doc,
            ['贡献人', '贡献类型', '贡献内容', '权重', '审核状态', '统计周期', '创建时间'],
            rows,
        )
    else:
        doc.add_paragraph('暂无贡献记录。')

    # ========== 页脚 ==========
    doc.add_paragraph()
    footer = doc.add_paragraph(f'报告生成时间：{timezone.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 写入 BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ============================================================
# 纯文本报告生成（python-docx 不可用时降级）
# ============================================================

def _generate_text_report(data):
    """生成纯文本报告，返回 BytesIO"""
    project = data['project']
    lines = []

    lines.append(f'项目报告：{project.name}')
    lines.append('=' * 60)
    lines.append('')

    # 一、项目基本信息
    lines.append('一、项目基本信息')
    lines.append('-' * 40)
    lines.append(f'项目名称：{project.name}')
    lines.append(f'项目编号：{project.code}')
    lines.append(f'项目负责人：{project.leader.name if project.leader else "未指定"}')
    lines.append(f'当前阶段：{project.get_current_stage_display()}')
    lines.append(f'项目状态：{project.get_status_display()}')
    lines.append(f'优先级：{project.get_priority_display()}')
    lines.append(f'开始时间：{_fmt_date(project.start_date) or "未指定"}')
    lines.append(f'预计结束：{_fmt_date(project.planned_end_date) or "未指定"}')
    lines.append(f'实际结束：{_fmt_date(project.actual_end_date) or "未指定"}')
    lines.append(f'项目简介：{project.intro or "无"}')
    lines.append(f'创建时间：{_fmt_datetime(project.created_at)}')
    lines.append('')

    # 二、阶段历程
    lines.append('二、阶段历程')
    lines.append('-' * 40)
    stage_logs = data['stage_logs']
    if stage_logs:
        for log in stage_logs:
            from_stage = log.get_from_stage_display() if log.from_stage else '初始'
            to_stage = log.get_to_stage_display()
            operator = log.operator.name if log.operator else '系统'
            lines.append(
                f'[{_fmt_datetime(log.created_at)}] {from_stage} -> {to_stage} '
                f'操作人：{operator} 备注：{log.note or "无"}'
            )
    else:
        lines.append('暂无阶段变更记录。')
    lines.append('')

    # 三、比赛记录
    lines.append('三、比赛记录')
    lines.append('-' * 40)
    competitions = data['competitions']
    if competitions:
        for c in competitions:
            lines.append(
                f'比赛名称：{c.name} | 级别：{c.get_level_display()} | '
                f'主办方：{c.organizer or "无"} | 状态：{c.get_status_display()} | '
                f'晋级：{"是" if c.is_promoted else "否"} | 获奖：{"是" if c.is_awarded else "否"} | '
                f'获奖等级：{c.award_level or "无"} | 结果公布：{_fmt_date(c.result_date) or "无"}'
            )
    else:
        lines.append('暂无比赛记录。')
    lines.append('')

    # 四、经费统计
    lines.append('四、经费统计')
    lines.append('-' * 40)
    budget = data['budget']
    if budget:
        lines.append(f'奖金总额：{budget.bonus_amount} 元')
        lines.append(f'其他收入：{budget.other_income} 元')
        lines.append(f'总收入：{budget.total_income} 元')
        lines.append(f'已用金额：{budget.used_amount} 元')
        lines.append(f'待报销：{budget.pending_reimbursement} 元')
        lines.append(f'剩余金额：{budget.remaining_amount} 元')
        lines.append(f'经费状态：{budget.get_status_display()}')
        lines.append(f'统计周期：{budget.period or "无"}')
    else:
        lines.append('暂无经费预算信息。')
    lines.append('')
    lines.append('经费明细：')
    expenses = data['expenses']
    if expenses:
        for e in expenses:
            lines.append(
                f'  - {e.title} | 金额：{e.amount}元 | 经办人：{e.spender.name if e.spender else "无"} | '
                f'日期：{_fmt_date(e.expense_date)} | 类别：{e.get_category_display()} | 用途：{e.purpose or "无"}'
            )
    else:
        lines.append('  暂无经费明细。')
    lines.append('')

    # 五、成员列表
    lines.append('五、成员列表')
    lines.append('-' * 40)
    members = data['members']
    if members:
        for m in members:
            lines.append(
                f'  - {m.user.name} | 邮箱：{m.user.email} | 角色：{m.get_role_in_project_display()} | '
                f'加入时间：{_fmt_datetime(m.joined_at)}'
            )
    else:
        lines.append('暂无项目成员。')
    lines.append('')

    # 六、知识产权
    lines.append('六、知识产权')
    lines.append('-' * 40)
    ip_apps = data['ip_apps']
    if ip_apps:
        for a in ip_apps:
            lines.append(
                f'  - {a.title} | 编号：{a.application_code} | 类型：{a.get_ip_type_display()} | '
                f'状态：{a.get_status_display()} | 主导撰写人：{a.main_writer.name if a.main_writer else "未指定"} | '
                f'提交：{_fmt_date(a.submit_date) or "未提交"} | 受理：{_fmt_date(a.accepted_date) or "未受理"} | '
                f'授权：{_fmt_date(a.authorized_date) or "未授权"}'
            )
    else:
        lines.append('暂无知识产权申请。')
    lines.append('')

    # 七、贡献汇总
    lines.append('七、贡献汇总')
    lines.append('-' * 40)
    contributions = data['contributions']
    if contributions:
        for c in contributions:
            lines.append(
                f'  - {c.user.name if c.user else "未知"} | 类型：{c.get_contribution_type_display()} | '
                f'内容：{c.content or c.description or "无"} | 权重：{c.weight} | '
                f'状态：{c.get_status_display()} | 周期：{c.period or "无"} | 创建：{_fmt_datetime(c.created_at)}'
            )
    else:
        lines.append('暂无贡献记录。')
    lines.append('')

    # 页脚
    lines.append('=' * 60)
    lines.append(f'报告生成时间：{timezone.now().strftime("%Y-%m-%d %H:%M")}')

    output = io.BytesIO()
    output.write('\n'.join(lines).encode('utf-8'))
    output.seek(0)
    return output


# ============================================================
# 对外接口
# ============================================================

def generate_project_report(project_id):
    """
    生成项目完整报告（Word 格式，使用 python-docx）
    若 python-docx 未安装，降级为纯文本导出
    报告包含：项目基本信息、阶段历程、比赛记录、经费统计、成员列表、知识产权、贡献汇总

    :param project_id: 项目ID
    :return: io.BytesIO 对象（含 docx 或 txt 内容）
    :raises ValueError: 项目不存在
    """
    data = _gather_project_data(project_id)

    if HAS_PYTHON_DOCX:
        return _generate_docx_report(data)
    else:
        return _generate_text_report(data)
